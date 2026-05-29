import logging
import os
from docx import Document as DocxDocument
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def _iter_table_text(table) -> list[str]:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = "\n".join(
                paragraph.text.strip()
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )
            if cell_text:
                cells.append(cell_text)
        if cells:
            rows.append(" | ".join(cells))
    return rows


def load_docx(file_path: str) -> list[Document]:
    source = os.path.basename(file_path)
    doc = DocxDocument(file_path)
    blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        blocks.extend(_iter_table_text(table))

    text = "\n".join(blocks)
    logger.info("[DOCX] %s: %d chars", source, len(text))
    return [Document(
        page_content=text,
        metadata={"source": source, "page": 1, "file_type": "docx"},
    )]
