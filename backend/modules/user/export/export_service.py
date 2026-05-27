# backend/modules/export/export_service.py
import json
import io
import logging
from backend.extensions import db
from backend.models.document import Document

logger = logging.getLogger(__name__)


class ExportService:
    STATUS_LABELS = {
        "da_duyet": "Đã duyệt",
        "cho_duyet": "Chờ duyệt",
        "yeu_cau_sua": "Yêu cầu sửa",
        "ban_nhap": "Bản nháp",
        "Đã duyệt": "Đã duyệt",
        "Chờ duyệt": "Chờ duyệt",
        "Bản nháp": "Bản nháp",
    }

    STATUS_NORMALIZE = {
        "Đã duyệt": "da_duyet", "da_duyet": "da_duyet",
        "Chờ duyệt": "cho_duyet", "cho_duyet": "cho_duyet",
        "Yêu cầu sửa": "yeu_cau_sua", "yeu_cau_sua": "yeu_cau_sua",
        "Bị từ chối": "yeu_cau_sua",
        "Bản nháp": "ban_nhap", "ban_nhap": "ban_nhap",
    }

    @staticmethod
    def _parse_doc(doc):
        so_hieu = "Chưa có"
        try:
            if doc.content:
                c = json.loads(doc.content)
                if isinstance(c, dict):
                    so_hieu = c.get("so_ky_hieu", "Chưa có")
        except Exception:
            pass

        raw_status = doc.status or "ban_nhap"
        status = ExportService.STATUS_NORMALIZE.get(raw_status, "ban_nhap")

        return {
            "id": doc.id,
            "so_hieu": so_hieu,
            "tieu_de": doc.title,
            "loai_van_ban": doc.doc_type,
            "trang_thai": status,
            "ngay_tao": doc.created_at.strftime("%d/%m/%Y"),
        }

    @staticmethod
    def get_exportable_documents(user_id):
        """Lấy danh sách văn bản có thể xuất (tất cả, hiển thị trạng thái)."""
        docs = (
            Document.query
            .filter_by(author_id=user_id)
            .order_by(Document.created_at.desc())
            .all()
        )
        return [ExportService._parse_doc(d) for d in docs]

    @staticmethod
    def export_documents(user_id, doc_ids, fmt="pdf"):
        """Xuất văn bản theo danh sách ID. Trả về (bytes, filename, mimetype)."""
        docs = (
            Document.query
            .filter(Document.id.in_(doc_ids), Document.author_id == user_id)
            .all()
        )
        if not docs:
            raise Exception("Không tìm thấy văn bản để xuất")

        if fmt == "docx":
            return ExportService._export_docx(docs)
        elif fmt == "odt":
            return ExportService._export_odt(docs)
        else:
            return ExportService._export_pdf(docs)

    @staticmethod
    def _export_pdf(docs):
        """Tạo PDF đơn giản chứa nội dung các văn bản."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas as pdf_canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            buf = io.BytesIO()
            c = pdf_canvas.Canvas(buf, pagesize=A4)
            w, h = A4

            for doc in docs:
                c.setFont("Helvetica-Bold", 14)
                c.drawString(72, h - 72, doc.title or "Không có tiêu đề")
                c.setFont("Helvetica", 10)
                c.drawString(72, h - 92, f"Loại: {doc.doc_type}  |  Trạng thái: {doc.status}")
                y = h - 120

                # Nội dung
                content_text = ""
                try:
                    if doc.content:
                        cj = json.loads(doc.content)
                        if isinstance(cj, dict):
                            content_text = cj.get("noi_dung", str(cj))
                        else:
                            content_text = str(doc.content)
                    else:
                        content_text = "(Không có nội dung)"
                except Exception:
                    content_text = str(doc.content or "")

                c.setFont("Helvetica", 11)
                for line in content_text.split("\n"):
                    if y < 72:
                        c.showPage()
                        y = h - 72
                    c.drawString(72, y, line[:90])
                    y -= 14

                c.showPage()

            c.save()
            buf.seek(0)
            fname = f"xuat_van_ban_{len(docs)}.pdf"
            return buf.getvalue(), fname, "application/pdf"
        except ImportError:
            # Fallback: trả về plain text nếu không có reportlab
            return ExportService._export_txt(docs, "pdf")

    @staticmethod
    def _export_docx(docs):
        """Tạo file DOCX."""
        try:
            from docx import Document as DocxDocument

            docx_doc = DocxDocument()
            for i, doc in enumerate(docs):
                if i > 0:
                    docx_doc.add_page_break()
                docx_doc.add_heading(doc.title or "Không có tiêu đề", level=1)
                docx_doc.add_paragraph(f"Loại: {doc.doc_type}  |  Trạng thái: {doc.status}")

                content_text = ""
                try:
                    if doc.content:
                        cj = json.loads(doc.content)
                        if isinstance(cj, dict):
                            content_text = cj.get("noi_dung", json.dumps(cj, ensure_ascii=False, indent=2))
                        else:
                            content_text = str(doc.content)
                    else:
                        content_text = "(Không có nội dung)"
                except Exception:
                    content_text = str(doc.content or "")

                docx_doc.add_paragraph(content_text)

            buf = io.BytesIO()
            docx_doc.save(buf)
            buf.seek(0)
            fname = f"xuat_van_ban_{len(docs)}.docx"
            return buf.getvalue(), fname, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except ImportError:
            return ExportService._export_txt(docs, "docx")

    @staticmethod
    def _export_odt(docs):
        """Fallback ODT → txt."""
        return ExportService._export_txt(docs, "odt")

    @staticmethod
    def _export_txt(docs, label="txt"):
        """Fallback text export."""
        lines = []
        for doc in docs:
            lines.append(f"=== {doc.title} ===")
            lines.append(f"Loại: {doc.doc_type}  |  Trạng thái: {doc.status}")
            lines.append("")
            try:
                if doc.content:
                    cj = json.loads(doc.content)
                    if isinstance(cj, dict):
                        lines.append(cj.get("noi_dung", json.dumps(cj, ensure_ascii=False, indent=2)))
                    else:
                        lines.append(str(doc.content))
                else:
                    lines.append("(Không có nội dung)")
            except Exception:
                lines.append(str(doc.content or ""))
            lines.append("\n")

        content = "\n".join(lines).encode("utf-8")
        fname = f"xuat_van_ban_{len(docs)}.txt"
        return content, fname, "text/plain; charset=utf-8"
